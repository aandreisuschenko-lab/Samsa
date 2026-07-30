"""Отчёт проверок и недельный дайджест в формате Word (.docx).

python-docx — мягкая зависимость: если библиотеки нет, пайплайн не падает,
а откатывается на markdown и печатает подсказку по установке.
"""
from __future__ import annotations

from collections import defaultdict
from datetime import date, timedelta
from pathlib import Path

from .checks import P1, P2, P3, Issue
from .deadlines import STAGE_LABELS, deadline_diff
from .loader import PlanData
from .report import _plural

try:  # мягкая зависимость
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    DOCX_AVAILABLE = False

INK = "0B1020"
MUTED = "5B6478"
BLUE = "1428A0"
RED = "C3271A"
AMBER = "B26B00"
GREEN = "147A4B"

PRIORITY_COLOR = {P1: RED, P2: AMBER, P3: MUTED}
PRIORITY_LABEL = {
    P1: "P1 · блокер",
    P2: "P2 · риск",
    P3: "P3 · дыра в покрытии",
}
PRIORITY_MEANING = {
    P1: "запуск под угрозой срыва, нужна реакция сегодня",
    P2: "риск не успеть или потерять эффективность",
    P3: "недоработка плана: дыры в покрытии, дисбаланс",
}


def _fmt(d: date | None) -> str:
    return d.strftime("%d.%m.%Y") if d else "—"


def _short(d: date | None) -> str:
    return d.strftime("%d.%m") if d else "—"


# --- вспомогательное оформление -------------------------------------------

def _shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _run(paragraph, text: str, *, bold=False, size=10, color=INK, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Calibri"
    return run


def _para(doc, text="", *, bold=False, size=10, color=INK, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        _run(p, text, bold=bold, size=size, color=color, italic=italic)
    return p


def _title(doc, text: str, size=18):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, text, bold=True, size=size, color=INK)
    return p


def _heading(doc, text: str, color=INK, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    _run(p, text, bold=True, size=size, color=color)
    return p


def _table(doc, headers: list[str], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    from docx.shared import Cm

    for i, cell in enumerate(table.rows[0].cells):
        cell.width = Cm(widths[i])
        _shade(cell, "EDF0F7")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        _run(p, headers[i], bold=True, size=9, color=INK)
    return table


def _row(table, values: list[str], widths: list[int], *, bold_first=False, color=INK):
    from docx.shared import Cm

    cells = table.add_row().cells
    for i, v in enumerate(values):
        cells[i].width = Cm(widths[i])
        p = cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        _run(p, v, bold=(bold_first and i == 0), size=9, color=color)
    return cells


# --- отчёт проверок --------------------------------------------------------

def build_checks_docx(plan: PlanData, issues: list[Issue], today: date, out_path: str | Path) -> Path:
    out_path = Path(out_path)
    doc = Document()

    section = doc.sections[0]
    from docx.shared import Cm

    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    _title(doc, "Отчёт проверок")
    _para(doc, plan.title, size=11, color=MUTED)
    _para(
        doc,
        f"Дата расчёта: {_fmt(today)}   ·   кампаний в плане: {len(plan.campaigns)}   ·   "
        f"найдено проблем: {len(issues)}",
        size=10, color=MUTED, space_after=10,
    )

    _heading(doc, "Сводка по приоритетам")
    widths = [4.0, 2.0, 11.0]
    table = _table(doc, ["Приоритет", "Кол-во", "Что это значит"], widths)
    for p in (P1, P2, P3):
        n = sum(1 for i in issues if i.priority == p)
        cells = _row(table, [PRIORITY_LABEL[p], str(n), PRIORITY_MEANING[p]], widths, bold_first=True)
        para = cells[0].paragraphs[0]
        para.runs[0].font.color.rgb = RGBColor.from_string(PRIORITY_COLOR[p])

    for p in (P1, P2, P3):
        block = [i for i in issues if i.priority == p]
        if not block:
            continue
        _heading(doc, PRIORITY_LABEL[p], color=PRIORITY_COLOR[p])
        for issue in block:
            para = _para(doc, space_after=2)
            _run(para, issue.title, bold=True, size=10)
            _para(doc, issue.detail, size=9, color=MUTED, space_after=2)
            if issue.action:
                para = _para(doc, space_after=8)
                _run(para, "Что делать: ", bold=True, size=9, color=BLUE)
                _run(para, issue.action, size=9)

    doc.add_page_break()
    _heading(doc, "Пересчёт дедлайнов: было / стало")
    _para(
        doc,
        "В исходном файле дедлайны посчитаны единым правилом −14 / −7 / −3 для всех кампаний. "
        "Ниже — расчёт по правилам вкладки «2. Правила дедлайнов» с переносом выходных "
        "на предыдущий рабочий день. Изменённые значения выделены цветом.",
        size=9, color=MUTED, space_after=8,
    )

    w = [1.8, 3.0, 1.8, 3.1, 3.1, 3.1]
    table = _table(doc, ["ID", "Тип кампании", "Старт", "Бриф", "Креатив", "Согласование"], w)
    for c in plan.campaigns:
        diffs = {d["stage"]: d for d in deadline_diff(c)}
        values = [c.id, c.campaign_type, _short(c.start)]
        colors = [INK, INK, INK]
        for stage in ("brief", "creative", "approval"):
            if stage in diffs:
                d = diffs[stage]
                values.append(f"{_short(d['was'])} → {_short(d['now'])} ({d['delta_days']:+d})")
                colors.append(RED if d["delta_days"] < 0 else AMBER)
            else:
                values.append(f"{_short(c.corrected.get(stage))} без изменений")
                colors.append(MUTED)
        cells = table.add_row().cells
        for i, v in enumerate(values):
            cells[i].width = Cm(w[i])
            para = cells[i].paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            _run(para, v, bold=(i == 0), size=8, color=colors[i])

    doc.save(out_path)
    return out_path


# --- недельный дайджест ----------------------------------------------------

def build_digest_docx(plan: PlanData, issues: list[Issue], today: date, out_path: str | Path) -> Path:
    out_path = Path(out_path)
    horizon = today + timedelta(days=7)
    burning = [i for i in issues if i.priority == P1]
    gaps = [i for i in issues if i.priority == P3]

    tasks = []
    for c in plan.campaigns:
        if c.is_done:
            continue
        for stage, d in sorted(c.corrected.items(), key=lambda x: x[1]):
            if stage == "upload":
                continue
            if d < today:
                tasks.append((d, c, stage, "просрочено"))
            elif d <= horizon:
                tasks.append((d, c, stage, "на этой неделе"))
    tasks.sort(key=lambda x: x[0])

    doc = Document()
    from docx.shared import Cm

    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    week_no = today.isocalendar().week
    _title(doc, f"Недельный дайджест · W{week_no}")
    _para(doc, f"Расчёт на {_fmt(today)}", size=10, color=MUTED)
    para = _para(doc, space_after=12)
    _run(para, _plural(len(burning), "блокер", "блокера", "блокеров"), bold=True, size=11, color=RED)
    _run(para, "   ·   ", size=11, color=MUTED)
    _run(para, _plural(len(tasks), "задача", "задачи", "задач") + " на неделю", bold=True, size=11, color=AMBER)
    _run(para, "   ·   ", size=11, color=MUTED)
    _run(para, _plural(len(gaps), "дыра", "дыры", "дыр") + " в плане", bold=True, size=11, color=MUTED)

    _heading(doc, "Что горит", color=RED)
    if burning:
        for i in burning[:8]:
            para = _para(doc, space_after=4)
            _run(para, "• ", size=10, color=RED)
            _run(para, i.title, bold=True, size=10)
            _run(para, " — " + (i.action or i.detail), size=10, color=MUTED)
        if len(burning) > 8:
            _para(doc, f"…и ещё {len(burning) - 8} — полный список в отчёте проверок",
                  size=9, color=MUTED, italic=True)
    else:
        _para(doc, "Блокеров нет, план идёт по графику.", size=10, color=GREEN)

    _heading(doc, "Что сделать на этой неделе", color=AMBER)
    if tasks:
        w = [2.6, 2.0, 3.0, 4.4, 3.9]
        table = _table(doc, ["Дедлайн", "Кампания", "Этап", "Ответственный", "Статус"], w)
        for d, c, stage, mark in tasks[:15]:
            late = mark == "просрочено"
            label = f"{_short(d)}{'  (просрочено)' if late else ''}"
            cells = _row(table, [label, c.id, STAGE_LABELS[stage], c.owner or "—", c.status], w,
                         bold_first=True)
            if late:
                cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(RED)
        if len(tasks) > 15:
            _para(doc, f"…и ещё {len(tasks) - 15} задач", size=9, color=MUTED, italic=True)
    else:
        _para(doc, "Дедлайнов в ближайшие 7 дней нет.", size=10, color=GREEN)

    _heading(doc, "Где дыры", color=MUTED)
    if gaps:
        for i in gaps:
            para = _para(doc, space_after=4)
            _run(para, "• ", size=10, color=MUTED)
            _run(para, i.title, bold=True, size=10)
            _run(para, " — " + (i.action or i.detail), size=10, color=MUTED)
    else:
        _para(doc, "Покрытие моделей, каналов и дистрибьюторов закрыто.", size=10, color=GREEN)

    _heading(doc, "Загрузка по неделям")
    by_week: dict[str, int] = defaultdict(int)
    for c in plan.campaigns:
        by_week[c.week or c.start.strftime("W%V")] += 1
    para = _para(doc)
    for i, (wk, n) in enumerate(sorted(by_week.items())):
        over = n > plan.thresholds.max_launches_per_week
        if i:
            _run(para, "   ·   ", size=10, color=MUTED)
        _run(para, f"{wk}: {n}" + (" (перегруз)" if over else ""), size=10,
             bold=over, color=RED if over else INK)

    doc.save(out_path)
    return out_path
